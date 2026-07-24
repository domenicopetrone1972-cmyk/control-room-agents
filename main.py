import base64
import io
import os
import re
import time
import random

import requests
import streamlit as st
from dotenv import load_dotenv
from crewai import Agent, Crew, Process, Task, LLM

load_dotenv()

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pypdf import PdfReader
from PIL import Image as PILImage

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, Image as RLImage,
)

# ------------------------------------------------------------------
# CONFIGURAZIONE PAGINA
# ------------------------------------------------------------------
st.set_page_config(page_title="Control Room · Multi-Agent", layout="wide", page_icon="◈")

CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;700&family=Inter:wght@400;500;600&display=swap');

:root {
    --bg: #0D1117;
    --panel: #161B22;
    --panel-border: #262E3A;
    --text: #E6E8EB;
    --text-muted: #8B93A7;
    --accent: #FF9F43;
    --accent-dim: #B96B22;
    --ok: #4ADE80;
    --err: #F87171;
}

html, body, [class*="css"] { font-family: 'Inter', sans-serif; color: var(--text); }

.stApp {
    background: var(--bg);
    background-image:
        radial-gradient(circle at 15% 0%, rgba(255, 159, 67, 0.06), transparent 35%),
        radial-gradient(circle at 85% 20%, rgba(255, 159, 67, 0.04), transparent 30%);
}

.control-header {
    display: flex; align-items: center; gap: 14px;
    padding: 4px 0 18px 0; border-bottom: 1px solid var(--panel-border); margin-bottom: 28px;
}
.control-header .dot {
    width: 10px; height: 10px; border-radius: 50%; background: var(--accent);
    animation: pulse 2.2s infinite; flex-shrink: 0;
}
@keyframes pulse {
    0% { box-shadow: 0 0 0 0 rgba(255, 159, 67, 0.55); }
    70% { box-shadow: 0 0 0 9px rgba(255, 159, 67, 0); }
    100% { box-shadow: 0 0 0 0 rgba(255, 159, 67, 0); }
}
.control-header h1 {
    font-family: 'JetBrains Mono', monospace; font-size: 1.5rem; font-weight: 700;
    letter-spacing: -0.02em; margin: 0; color: var(--text);
}
.control-header .subtitle {
    font-family: 'JetBrains Mono', monospace; font-size: 0.78rem;
    color: var(--text-muted); letter-spacing: 0.04em; margin-top: 2px;
}

.pipeline {
    display: flex; align-items: stretch; gap: 0; margin: 6px 0 30px 0;
    border: 1px solid var(--panel-border); border-radius: 10px; overflow: hidden; background: var(--panel);
    flex-wrap: wrap;
}
.pipeline .stage { flex: 1; min-width: 200px; padding: 16px 18px; position: relative; }
.pipeline .stage:not(:last-child) { border-right: 1px dashed var(--panel-border); }
.pipeline .stage .tag {
    font-family: 'JetBrains Mono', monospace; font-size: 0.68rem; color: var(--accent);
    letter-spacing: 0.08em; text-transform: uppercase;
}
.pipeline .stage .role { font-weight: 600; font-size: 0.95rem; margin-top: 4px; color: var(--text); }
.pipeline .stage .desc { font-size: 0.8rem; color: var(--text-muted); margin-top: 3px; line-height: 1.35; }
.pipeline .arrow {
    display: flex; align-items: center; justify-content: center; padding: 0 4px;
    color: var(--accent-dim); font-family: 'JetBrains Mono', monospace; font-size: 1rem;
}

.section-label {
    font-family: 'JetBrains Mono', monospace; font-size: 0.72rem; color: var(--text-muted);
    letter-spacing: 0.1em; text-transform: uppercase; margin-bottom: 6px;
}

section[data-testid="stSidebar"] { background: var(--panel); border-right: 1px solid var(--panel-border); }
section[data-testid="stSidebar"] h2 {
    font-family: 'JetBrains Mono', monospace; font-size: 0.85rem; color: var(--accent);
    letter-spacing: 0.08em; text-transform: uppercase;
}

.stTextArea textarea, .stTextInput input {
    background: var(--panel) !important; border: 1px solid var(--panel-border) !important;
    color: var(--text) !important; border-radius: 8px !important; font-family: 'Inter', sans-serif;
}
.stTextArea textarea:focus, .stTextInput input:focus {
    border-color: var(--accent) !important; box-shadow: 0 0 0 1px var(--accent) !important;
}
.stSelectbox div[data-baseweb="select"] > div {
    background: var(--panel) !important; border: 1px solid var(--panel-border) !important; border-radius: 8px !important;
}

section[data-testid="stFileUploaderDropzone"] {
    background: var(--panel) !important; border: 1px dashed var(--panel-border) !important; border-radius: 8px !important;
}

.stButton button {
    background: var(--accent) !important; color: #14100A !important; border: none !important;
    border-radius: 8px !important; font-weight: 600 !important; letter-spacing: 0.02em;
    padding: 0.55rem 1.4rem !important; transition: transform 0.12s ease, box-shadow 0.12s ease;
}
.stButton button:hover { box-shadow: 0 0 0 3px rgba(255, 159, 67, 0.25); transform: translateY(-1px); }

.stDownloadButton button {
    background: var(--panel) !important; color: var(--text) !important;
    border: 1px solid var(--accent-dim) !important; border-radius: 8px !important; font-weight: 500 !important;
}
.stDownloadButton button:hover { border-color: var(--accent) !important; color: var(--accent) !important; }

.result-panel {
    background: var(--panel); border: 1px solid var(--panel-border); border-left: 3px solid var(--ok);
    border-radius: 8px; padding: 20px 22px; margin-top: 12px;
}

div[data-testid="stAlert"] { border-radius: 8px; font-family: 'Inter', sans-serif; }
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# ------------------------------------------------------------------
# DEFINIZIONE DEL TEAM DI AGENTI (istruzioni tarate per output sintetico)
# ------------------------------------------------------------------
AGENTS_CONFIG = [
    {
        "key": "researcher",
        "tag": "Agente 01",
        "role": "Ricercatore Senior",
        "goal": "Raccogliere solo i dati più rilevanti sul task richiesto, in modo sintetico.",
        "backstory": "Sei un analista esperto che va dritto al punto, senza divagazioni.",
        "desc": "Raccoglie i dati essenziali sul task.",
        "task": (lambda task_input: f"Fai una ricerca sintetica su questo argomento: {task_input}. "
                                     f"Individua solo i 3-4 dati o trend più rilevanti. Massimo 100 parole."),
        "expected_output": "Elenco puntato con i 3-4 punti chiave più rilevanti (massimo 100 parole).",
    },
    {
        "key": "analyst",
        "tag": "Agente 02",
        "role": "Analista Critico",
        "goal": "Individuare rapidamente i rischi e le opportunità più significativi.",
        "backstory": "Sei un analista di mercato abituato a sintesi rapide per decisioni operative.",
        "desc": "Evidenzia i 2-3 rischi/opportunità più rilevanti.",
        "task": (lambda task_input: "Analizza i dati raccolti dal ricercatore. Individua solo i 2-3 elementi più "
                                     "critici tra rischi e opportunità. Massimo 80 parole, elenco puntato."),
        "expected_output": "Elenco puntato con i 2-3 rischi/opportunità più rilevanti (massimo 80 parole).",
    },
    {
        "key": "writer",
        "tag": "Agente 03",
        "role": "Redattore di Report",
        "goal": "Scrivere un report breve, diretto e professionale.",
        "backstory": "Sei un copywriter che scrive per manager che hanno 30 secondi da dedicare alla lettura.",
        "desc": "Scrive un report breve in Markdown.",
        "task": (lambda task_input: "Scrivi un report in italiano basato sui dati e sull'analisi. Massimo 200 "
                                     "parole totali. Markdown con 2-3 titoli (##) ed elenchi puntati brevi. "
                                     "Nessuna frase di riempimento."),
        "expected_output": "Report in Markdown, massimo 200 parole, con 2-3 titoli ed elenchi puntati essenziali.",
    },
    {
        "key": "editor",
        "tag": "Agente 04",
        "role": "Editor & QA",
        "goal": "Tagliare tutto il superfluo e restituire un report pronto per un cliente.",
        "backstory": "Sei un editor spietato con le ripetizioni: se una frase non aggiunge valore, la elimini.",
        "desc": "Taglia il superfluo, output finale pronto.",
        "task": (lambda task_input: "Rileggi il report e riducilo se necessario: non deve superare le 200 parole "
                                     "totali. Elimina ripetizioni e dettagli superflui. Output pulito, in Markdown, "
                                     "senza commenti fuori dal report."),
        "expected_output": "Report finale conciso (massimo 200 parole), in Markdown pulito.",
    },
]

# ------------------------------------------------------------------
# HEADER
# ------------------------------------------------------------------
st.markdown(
    """
    <div class="control-header">
        <div class="dot"></div>
        <div>
            <h1>CONTROL ROOM</h1>
            <div class="subtitle">MULTI-AGENT ORCHESTRATION · LOCAL WORKSPACE</div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

# ------------------------------------------------------------------
# SIDEBAR
# ------------------------------------------------------------------
st.sidebar.header("Configurazione")
default_api_key = os.getenv("GEMINI_API_KEY", "")
api_key = st.sidebar.text_input(
    "API Key",
    type="password",
    value=default_api_key,
    placeholder="Incolla qui la tua chiave",
    help="Caricata automaticamente dal file .env. Puoi sovrascriverla qui se serve.",
)

MODEL_OPTIONS = {
    "Gemini 3.5 Flash (consigliato)": "gemini/gemini-3.5-flash",
    "Gemini 3.5 Flash-Lite (più economico)": "gemini/gemini-3.5-flash-lite",
    "Gemini 2.5 Flash (alternativa stabile)": "gemini/gemini-2.5-flash",
}
model_label = st.sidebar.selectbox("Modello Gemini", list(MODEL_OPTIONS.keys()))
selected_model = MODEL_OPTIONS[model_label]

generate_image = st.sidebar.checkbox("Genera immagine di copertina", value=True)

st.sidebar.markdown("---")
st.sidebar.markdown(
    f"<span style='font-family:JetBrains Mono,monospace; font-size:0.72rem; color:#8B93A7;'>"
    f"STATO<br>{len(AGENTS_CONFIG)} agenti · processo sequenziale</span>",
    unsafe_allow_html=True,
)

# ------------------------------------------------------------------
# PIPELINE VISUAL
# ------------------------------------------------------------------
stages_html = ""
for i, cfg in enumerate(AGENTS_CONFIG):
    stages_html += f"""
        <div class="stage">
            <div class="tag">{cfg['tag']}</div>
            <div class="role">{cfg['role']}</div>
            <div class="desc">{cfg['desc']}</div>
        </div>
    """
    if i < len(AGENTS_CONFIG) - 1:
        stages_html += '<div class="arrow">→</div>'

st.markdown(f'<div class="pipeline">{stages_html}</div>', unsafe_allow_html=True)

# ------------------------------------------------------------------
# AREA PRINCIPALE — task + allegati
# ------------------------------------------------------------------
st.markdown('<div class="section-label">Compito da assegnare al team</div>', unsafe_allow_html=True)
task_input = st.text_area(
    "task_input",
    "Analizza le tendenze di mercato per il lancio di un software gestionale.",
    label_visibility="collapsed",
    height=100,
)

st.markdown('<div class="section-label" style="margin-top:16px;">Allegati (opzionale)</div>', unsafe_allow_html=True)
uploaded_files = st.file_uploader(
    "uploaded_files",
    type=["txt", "pdf", "docx", "png", "jpg", "jpeg"],
    accept_multiple_files=True,
    label_visibility="collapsed",
    help="Documenti (txt/pdf/docx) vengono letti come contesto. Le immagini vengono descritte automaticamente.",
)

run_clicked = st.button("▶  Avvia il Team di Agenti")


def kickoff_with_retry(crew, max_retries=4):
    """Esegue il crew con retry a backoff esponenziale in caso di errore 429 (quota superata)."""
    for attempt in range(max_retries):
        try:
            return crew.kickoff()
        except Exception as e:
            error_str = str(e)
            is_rate_limit = "429" in error_str or "RESOURCE_EXHAUSTED" in error_str or "quota" in error_str.lower()
            if is_rate_limit and attempt < max_retries - 1:
                wait_time = (2 ** attempt) + random.uniform(0, 1)
                st.warning(
                    f"Quota API raggiunta (429). Riprovo tra {wait_time:.1f}s "
                    f"(tentativo {attempt + 1}/{max_retries})..."
                )
                time.sleep(wait_time)
                continue
            raise
    raise RuntimeError("Numero massimo di tentativi raggiunto dopo errori 429 ripetuti.")


# ------------------------------------------------------------------
# LETTURA ALLEGATI
# ------------------------------------------------------------------
MAX_CHARS_PER_DOC = 4000


def extract_text_from_txt(file) -> str:
    raw = file.read()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="ignore")


def extract_text_from_pdf(file) -> str:
    reader = PdfReader(file)
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return text


def extract_text_from_docx_file(file) -> str:
    document = Document(file)
    return "\n".join(p.text for p in document.paragraphs)


def describe_image_with_gemini(api_key: str, image_bytes: bytes, mime_type: str) -> str:
    """Usa Gemini (vision) per descrivere brevemente un'immagine allegata."""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent?key={api_key}"
    payload = {
        "contents": [{
            "parts": [
                {"text": "Descrivi in italiano, in massimo 2 frasi, il contenuto di questa immagine, "
                          "evidenziando dettagli utili per un report di lavoro."},
                {"inlineData": {"mimeType": mime_type, "data": base64.b64encode(image_bytes).decode("utf-8")}},
            ]
        }]
    }
    resp = requests.post(url, json=payload, timeout=30)
    resp.raise_for_status()
    data = resp.json()
    return data["candidates"][0]["content"]["parts"][0]["text"].strip()


def build_attachments_context(files, api_key: str) -> str:
    """Estrae testo/descrizioni da tutti gli allegati e li unisce in un unico blocco di contesto."""
    blocks = []
    for f in files:
        ext = f.name.lower().rsplit(".", 1)[-1]
        try:
            if ext == "txt":
                content = extract_text_from_txt(f)[:MAX_CHARS_PER_DOC]
                blocks.append(f"--- Documento: {f.name} ---\n{content}")
            elif ext == "pdf":
                content = extract_text_from_pdf(f)[:MAX_CHARS_PER_DOC]
                blocks.append(f"--- Documento: {f.name} ---\n{content}")
            elif ext == "docx":
                content = extract_text_from_docx_file(f)[:MAX_CHARS_PER_DOC]
                blocks.append(f"--- Documento: {f.name} ---\n{content}")
            elif ext in ("png", "jpg", "jpeg"):
                mime = "image/png" if ext == "png" else "image/jpeg"
                description = describe_image_with_gemini(api_key, f.read(), mime)
                blocks.append(f"--- Immagine: {f.name} ---\n{description}")
        except Exception as e:
            st.warning(f"Non sono riuscito a leggere l'allegato '{f.name}': {e}")
    return "\n\n".join(blocks)


# ------------------------------------------------------------------
# GENERAZIONE IMMAGINE DI COPERTINA
# ------------------------------------------------------------------
def generate_cover_image(api_key: str, topic: str) -> bytes | None:
    """Genera un'immagine di copertina a tema con Gemini (Nano Banana)."""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-image:generateContent?key={api_key}"
    prompt = (
        f"Immagine editoriale professionale, stile minimal e pulito, senza testo scritto, "
        f"colori sobri, adatta come copertina di un report aziendale. Tema: {topic}"
    )
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"responseModalities": ["IMAGE", "TEXT"]},
    }
    resp = requests.post(url, json=payload, timeout=60)
    resp.raise_for_status()
    data = resp.json()
    for part in data["candidates"][0]["content"]["parts"]:
        inline = part.get("inlineData") or part.get("inline_data")
        if inline and inline.get("data"):
            return base64.b64decode(inline["data"])
    return None


# ------------------------------------------------------------------
# EXPORT: Markdown (+ immagine) -> DOCX / PDF professionali
# ------------------------------------------------------------------
def _split_bold(text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    segments = []
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            segments.append((part[2:-2], True))
        elif part:
            segments.append((part, False))
    return segments


def markdown_to_docx_bytes(md_text: str, title: str, image_bytes: bytes | None = None) -> bytes:
    doc = Document()

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if image_bytes:
        doc.add_picture(io.BytesIO(image_bytes), width=Inches(6))
        doc.add_paragraph("")

    for raw_line in md_text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            for text, bold in _split_bold(line.strip()[2:]):
                run = p.add_run(text)
                run.bold = bold
        else:
            p = doc.add_paragraph()
            for text, bold in _split_bold(line):
                run = p.add_run(text)
                run.bold = bold

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_pdf_bytes(md_text: str, title: str, image_bytes: bytes | None = None) -> bytes:
    buf = io.BytesIO()
    left_margin = right_margin = top_margin = bottom_margin = 2 * cm
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=left_margin, rightMargin=right_margin, topMargin=top_margin, bottomMargin=bottom_margin,
    )
    avail_width = A4[0] - left_margin - right_margin

    styles = getSampleStyleSheet()
    accent = HexColor("#B96B22")
    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"], textColor=accent, spaceAfter=18)
    h2_style = ParagraphStyle("H2Custom", parent=styles["Heading2"], textColor=accent, spaceBefore=14, spaceAfter=6)
    h3_style = ParagraphStyle("H3Custom", parent=styles["Heading3"], spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle("BodyCustom", parent=styles["BodyText"], leading=15, spaceAfter=6)

    def to_inline_markup(text):
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)

    elements = [Paragraph(title, title_style)]

    if image_bytes:
        try:
            img = PILImage.open(io.BytesIO(image_bytes))
            w, h = img.size
            target_width = avail_width
            target_height = target_width * (h / w)
            max_height = 10 * cm
            if target_height > max_height:
                target_height = max_height
                target_width = target_height * (w / h)
            elements.append(RLImage(io.BytesIO(image_bytes), width=target_width, height=target_height))
            elements.append(Spacer(1, 10))
        except Exception:
            pass

    elements.append(Spacer(1, 6))
    bullet_buffer = []

    def flush_bullets():
        if bullet_buffer:
            items = [ListItem(Paragraph(to_inline_markup(b), body_style)) for b in bullet_buffer]
            elements.append(ListFlowable(items, bulletType="bullet", leftIndent=14))
            bullet_buffer.clear()

    for raw_line in md_text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            flush_bullets()
            elements.append(Spacer(1, 4))
            continue
        if line.startswith("### "):
            flush_bullets()
            elements.append(Paragraph(to_inline_markup(line[4:].strip()), h3_style))
        elif line.startswith("## "):
            flush_bullets()
            elements.append(Paragraph(to_inline_markup(line[3:].strip()), h2_style))
        elif line.startswith("# "):
            flush_bullets()
            elements.append(Paragraph(to_inline_markup(line[2:].strip()), h2_style))
        elif line.strip().startswith(("- ", "* ")):
            bullet_buffer.append(line.strip()[2:])
        else:
            flush_bullets()
            elements.append(Paragraph(to_inline_markup(line), body_style))

    flush_bullets()
    doc.build(elements)
    return buf.getvalue()


def render_markdown_preview(md_text: str):
    html = md_text
    html = re.sub(r"^### (.*)$", r"<h4>\1</h4>", html, flags=re.MULTILINE)
    html = re.sub(r"^## (.*)$", r"<h3>\1</h3>", html, flags=re.MULTILINE)
    html = re.sub(r"^# (.*)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
    html = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", html)
    html = re.sub(r"^[-*] (.*)$", r"<li>\1</li>", html, flags=re.MULTILINE)
    html = html.replace("\n", "<br>")
    st.markdown(f'<div class="result-panel">{html}</div>', unsafe_allow_html=True)


# ------------------------------------------------------------------
# ESECUZIONE
# ------------------------------------------------------------------
if run_clicked:
    if not api_key:
        st.error("Per favore, inserisci prima la tua chiave API nella barra laterale.")
    elif not task_input:
        st.warning("Per favore, inserisci un compito da assegnare agli agenti.")
    else:
        # 0. Allegati -> contesto aggiuntivo
        combined_input = task_input
        if uploaded_files:
            with st.spinner("Leggo gli allegati..."):
                attachments_context = build_attachments_context(uploaded_files, api_key)
            if attachments_context:
                combined_input = f"{task_input}\n\nContesto aggiuntivo fornito dall'utente:\n{attachments_context}"

        with st.spinner("Gli agenti sono al lavoro..."):
            try:
                gemini_llm = LLM(model=selected_model, api_key=api_key)

                # 1. Crea gli agenti dalla configurazione
                agents = {}
                for cfg in AGENTS_CONFIG:
                    agents[cfg["key"]] = Agent(
                        role=cfg["role"],
                        goal=cfg["goal"],
                        backstory=cfg["backstory"],
                        verbose=True,
                        llm=gemini_llm,
                    )

                # 2. Crea i task in sequenza, incatenando il contesto
                tasks = []
                for cfg in AGENTS_CONFIG:
                    task = Task(
                        description=cfg["task"](combined_input),
                        expected_output=cfg["expected_output"],
                        agent=agents[cfg["key"]],
                        context=list(tasks) if tasks else None,
                    )
                    tasks.append(task)

                # 3. Crew
                my_crew = Crew(
                    agents=list(agents.values()),
                    tasks=tasks,
                    process=Process.sequential,
                    verbose=True,
                )

                # 4. Esecuzione con retry sui 429
                result = kickoff_with_retry(my_crew)
                report_text = result.raw if hasattr(result, "raw") else str(result)

                # 5. Immagine di copertina (facoltativa, non blocca il flusso se fallisce)
                cover_image_bytes = None
                if generate_image:
                    try:
                        with st.spinner("Genero immagine di copertina..."):
                            cover_image_bytes = generate_cover_image(api_key, task_input)
                    except Exception:
                        st.info("Non sono riuscito a generare l'immagine di copertina (facoltativa) — il report procede senza.")

                # 6. Risultato a schermo
                st.markdown('<div class="section-label" style="margin-top:24px;">Risultato</div>', unsafe_allow_html=True)
                if cover_image_bytes:
                    st.image(cover_image_bytes, use_container_width=True)
                render_markdown_preview(report_text)

                # 7. Export
                docx_bytes = markdown_to_docx_bytes(report_text, "Report — Control Room", cover_image_bytes)
                pdf_bytes = markdown_to_pdf_bytes(report_text, "Report — Control Room", cover_image_bytes)

                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "⬇  Scarica come Word (.docx)",
                        data=docx_bytes,
                        file_name="report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                with col2:
                    st.download_button(
                        "⬇  Scarica come PDF",
                        data=pdf_bytes,
                        file_name="report.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                    )

            except Exception as e:
                error_str = str(e)
                if "404" in error_str and "NOT_FOUND" in error_str:
                    st.error(
                        "Errore 404: il modello selezionato non è disponibile per la tua API key. "
                        "Prova a cambiare modello dalla barra laterale."
                    )
                elif "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                    st.error(
                        "Errore 429: hai superato la quota della tua API key "
                        "(troppe richieste in poco tempo). Aspetta qualche minuto, "
                        "oppure passa a un modello più economico come Flash-Lite."
                    )
                elif "API_KEY_INVALID" in error_str or "401" in error_str:
                    st.error("La API key non è valida. Controlla di averla copiata correttamente.")
                else:
                    st.error(f"Errore durante l'esecuzione: {e}")